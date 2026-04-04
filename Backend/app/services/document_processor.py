"""
Document processing service for extracting text from various file formats.
Supports: PDF, EPUB, TXT, DOCX, MOBI, AZW/AZW3
"""

import os
import shutil
import time
import uuid
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

# EPUB processing
import ebooklib
import pdfplumber
from bs4 import BeautifulSoup

# DOCX processing
from docx import Document as DocxDocument
from ebooklib import epub

# PDF processing
# import PyPDF2
from pypdf import PdfReader

# MOBI/AZW processing
try:
    import mobi
except ImportError:
    mobi = None

from fastapi import UploadFile
from loguru import logger

from app.config import get_settings
from app.models.schemas import (
    DocumentExtractionResult,
    DocumentMetadata,
    DocumentStatus,
    DocumentType,
)

settings = get_settings()


class DocumentProcessor:
    """
    Process and extract text from various document formats.
    """

    SUPPORTED_FORMATS = {
        "pdf": ["application/pdf"],
        "epub": ["application/epub+zip"],
        "txt": ["text/plain"],
        "docx": [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ],
        "mobi": ["application/x-mobipocket-ebook"],
        "azw": ["application/vnd.amazon.ebook"],
        "azw3": ["application/vnd.amazon.ebook"],
    }

    def __init__(self):
        """Initialize document processor."""
        self.upload_dir = settings.get_pdf_storage_path()
        self.max_file_size = settings.max_file_size_bytes

        # Ensure upload directory exists
        self.upload_dir.mkdir(parents=True, exist_ok=True)

        logger.info(f"DocumentProcessor initialized. Upload dir: {self.upload_dir}")

    def validate_file(self, file: UploadFile) -> Tuple[bool, Optional[str]]:
        """
        Validate uploaded file.

        Args:
            file: Uploaded file

        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check if file exists
        if not file or not file.filename:
            return False, "No file provided"
        filename = file.filename

        # Get file extension
        file_ext = self._get_file_extension(filename)

        # Check if extension is supported
        if file_ext not in self.SUPPORTED_FORMATS:
            supported = ", ".join(self.SUPPORTED_FORMATS.keys())
            return False, f"Unsupported file type. Supported: {supported}"

        # Check file size
        if file.size and file.size > self.max_file_size:
            max_mb = self.max_file_size / (1024 * 1024)
            return False, f"File size exceeds maximum of {max_mb}MB"

        return True, None

    def _get_file_extension(self, filename: str) -> str:
        """Get file extension in lowercase."""
        return Path(filename).suffix.lower().lstrip(".")

    def _get_file_type(self, filename: str) -> DocumentType:
        """Determine document type from filename."""
        ext = self._get_file_extension(filename)

        # Map extensions to DocumentType enum
        type_mapping = {
            "pdf": DocumentType.PDF,
            "epub": DocumentType.EPUB,
            "txt": DocumentType.TXT,
            "docx": DocumentType.DOCX,
            "mobi": DocumentType.MOBI,
            "azw": DocumentType.AZW,
            "azw3": DocumentType.AZW3,
        }

        return type_mapping.get(ext, DocumentType.PDF)

    async def save_file(self, file: UploadFile) -> Tuple[str, Path]:
        """
        Save uploaded file to disk.

        Args:
            file: Uploaded file

        Returns:
            Tuple of (document_id, file_path)
        """
        # Generate unique document ID
        document_id = str(uuid.uuid4())

        # Create filename with document ID
        filename = file.filename
        if not filename:
            raise ValueError("No filename provided")
        file_ext = self._get_file_extension(filename)
        safe_filename = f"{document_id}.{file_ext}"
        file_path = self.upload_dir / safe_filename

        # Save file
        try:
            content = await file.read()

            with open(file_path, "wb") as f:
                f.write(content)

            logger.info(f"File saved: {file_path} ({len(content)} bytes)")

            return document_id, file_path

        except Exception as e:
            logger.error(f"Failed to save file: {e}")
            raise

    def extract_text(
        self, file_path: Path, file_type: DocumentType
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from document based on file type.

        Args:
            file_path: Path to document file
            file_type: Type of document

        Returns:
            Tuple of (extracted_text, metadata_dict)
        """

        extractors = {
            DocumentType.PDF: self._extract_pdf,
            DocumentType.EPUB: self._extract_epub,
            DocumentType.TXT: self._extract_txt,
            DocumentType.DOCX: self._extract_docx,
            DocumentType.MOBI: self._extract_mobi,
            DocumentType.AZW: self._extract_mobi,
            DocumentType.AZW3: self._extract_mobi,
        }

        extractor = extractors.get(file_type)

        if not extractor:
            raise ValueError(f"No extractor available for {file_type}")

        try:
            return extractor(file_path)
        except Exception as e:
            logger.error(f"Text extraction failed for {file_type}: {e}")
            raise

    # def _extract_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
    #     """Extract text from PDF file."""
    #     logger.info(f"Extracting text from PDF: {file_path}")

    #     text_parts = []

    #     # Read the extracted
    #     metadata = {}

    #     try:
    #         # Try pdfplumber first (better for complex PDFs)
    #         with pdfplumber.open(file_path) as pdf:
    #             metadata["page_count"] = len(pdf.pages)

    #             for page in pdf.pages:
    #                 page_text = page.extract_text()
    #                 if page_text:
    #                     text_parts.append(page_text)

    #             # Extract metadata if available
    #             if pdf.metadata:
    #                 metadata["title"] = pdf.metadata.get("Title")
    #                 metadata["author"] = pdf.metadata.get("Author")

    #     except Exception as e:
    #         logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")

    #         # Fallback to PyPDF2
    #         try:
    #             with open(file_path, "rb") as f:
    #                 pdf_reader = PyPDF2.PdfReader(f)
    #                 metadata["page_count"] = len(pdf_reader.pages)

    #                 for page in pdf_reader.pages:
    #                     page_text = page.extract_text()
    #                     if page_text:
    #                         text_parts.append(page_text)

    #                 # Extract metadata
    #                 if pdf_reader.metadata:
    #                     metadata["title"] = pdf_reader.metadata.get("/Title")
    #                     metadata["author"] = pdf_reader.metadata.get("/Author")

    #         except Exception as e2:
    #             logger.error(f"PyPDF2 also failed: {e2}")
    #             raise

    #     full_text = "\n\n".join(text_parts)
    #     metadata["word_count"] = len(full_text.split())

    #     logger.info(f"PDF extraction complete: {metadata['word_count']} words")

    #     return full_text, metadata
    #
    def _extract_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file."""
        logger.info(f"Extracting text from PDF: {file_path}")

        text_parts = []
        metadata = {}

        try:
            # Try pdfplumber first (better for complex PDFs)
            with pdfplumber.open(file_path) as pdf:
                metadata["page_count"] = len(pdf.pages)

                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                # Extract metadata if available
                if pdf.metadata:
                    metadata["title"] = pdf.metadata.get("Title")
                    metadata["author"] = pdf.metadata.get("Author")

        except Exception as e:
            logger.warning(f"pdfplumber failed, trying pypdf: {e}")

            # Fallback to pypdf
            try:
                with open(file_path, "rb") as f:
                    pdf_reader = PdfReader(f)
                    metadata["page_count"] = len(pdf_reader.pages)

                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)

                    # Extract metadata
                    if pdf_reader.metadata:
                        metadata["title"] = pdf_reader.metadata.get("/Title")
                        metadata["author"] = pdf_reader.metadata.get("/Author")

            except Exception as e2:
                logger.error(f"pypdf also failed: {e2}")
                raise

        full_text = "\n\n".join(text_parts)
        metadata["word_count"] = len(full_text.split())

        logger.info(f"PDF extraction complete: {metadata['word_count']} words")

        return full_text, metadata

    def _extract_epub(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from EPUB file."""
        logger.info(f"Extracting text from EPUB: {file_path}")

        text_parts = []
        metadata = {}

        try:
            book = epub.read_epub(file_path)

            # Extract metadata
            metadata["title"] = book.get_metadata("DC", "title")
            metadata["author"] = book.get_metadata("DC", "creator")

            if metadata["title"]:
                metadata["title"] = (
                    metadata["title"][0][0] if metadata["title"] else None
                )
            if metadata["author"]:
                metadata["author"] = (
                    metadata["author"][0][0] if metadata["author"] else None
                )

            # Extract text from all items
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    content = item.get_content()
                    soup = BeautifulSoup(content, "html.parser")
                    text = soup.get_text(separator="\n", strip=True)
                    if text:
                        text_parts.append(text)

            full_text = "\n\n".join(text_parts)
            metadata["word_count"] = len(full_text.split())
            metadata["page_count"] = len(text_parts)  # Approximate

            logger.info(f"EPUB extraction complete: {metadata['word_count']} words")

            return full_text, metadata

        except Exception as e:
            logger.error(f"EPUB extraction failed: {e}")
            raise

    def _extract_txt(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from TXT file."""
        logger.info(f"Extracting text from TXT: {file_path}")

        metadata = {}

        try:
            # Try different encodings
            encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

            text = None
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        text = f.read()
                    logger.info(f"Successfully read TXT with encoding: {encoding}")
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                raise ValueError(
                    "Could not decode text file with any supported encoding"
                )

            metadata["word_count"] = len(text.split())
            metadata["page_count"] = 1

            logger.info(f"TXT extraction complete: {metadata['word_count']} words")

            return text, metadata

        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            raise

    def _extract_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file."""
        logger.info(f"Extracting text from DOCX: {file_path}")

        text_parts = []
        metadata = {}

        try:
            doc = DocxDocument(str(file_path))

            # Extract core properties
            core_props = doc.core_properties
            metadata["title"] = core_props.title
            metadata["author"] = core_props.author

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            full_text = "\n\n".join(text_parts)
            metadata["word_count"] = len(full_text.split())
            metadata["page_count"] = 1  # DOCX doesn't have clear page concept

            logger.info(f"DOCX extraction complete: {metadata['word_count']} words")

            return full_text, metadata

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise

    def _extract_mobi(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from MOBI/AZW/AZW3 file."""
        logger.info(f"Extracting text from MOBI/AZW: {file_path}")

        if mobi is None:
            raise ImportError(
                "mobi library not installed. Install with: pip install mobi"
            )

        metadata: Dict[str, Any] = {}
        tempdir = None

        try:
            # Extract MOBI contents
            tempdir, extracted_path = mobi.extract(str(file_path))

            # Locate extracted HTML (handle different MOBI versions)
            html_file = Path(tempdir) / "mobi7" / "book.html"
            if not html_file.exists():
                html_file = Path(tempdir) / "mobi8" / "book.html"

            # Fallback to returned extracted path if needed
            if not html_file.exists() and extracted_path:
                html_file = Path(extracted_path)

            if not html_file.exists():
                raise FileNotFoundError("Extracted HTML file not found")

            # Read and parse HTML
            with open(html_file, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            soup = BeautifulSoup(content, "html.parser")
            text = soup.get_text(separator="\n", strip=True)

            metadata["word_count"] = len(text.split())
            metadata["page_count"] = 1  # MOBI has no reliable page structure

            logger.info(f"MOBI extraction complete: {metadata['word_count']} words")

            return text, metadata

        except Exception as e:
            logger.error(f"MOBI extraction failed: {e}")
            # Fail immediately without fallback to raw decode
            raise

        finally:
            # Always clean up temporary directory
            if tempdir:
                shutil.rmtree(tempdir, ignore_errors=True)

    async def process_document(
        self, file: UploadFile, title: Optional[str] = None, tags: Optional[list] = None
    ) -> DocumentExtractionResult:
        """
        Process uploaded document: save, extract text, create metadata.

        Args:
            file: Uploaded file
            title: Optional custom title
            tags: Optional tags

        Returns:
            DocumentExtractionResult with extracted content and metadata
        """
        start_time = time.time()

        try:
            # Validate file
            is_valid, error_msg = self.validate_file(file)
            if not is_valid:
                raise ValueError(error_msg)

            # Save file
            document_id, file_path = await self.save_file(file)

            # Get file type
            filename = file.filename
            if not filename:
                raise ValueError("No filename provided")
            file_type = self._get_file_type(filename)

            # Extract text
            text_content, extracted_metadata = self.extract_text(file_path, file_type)

            # Create metadata
            metadata = DocumentMetadata(
                title=title or extracted_metadata.get("title") or filename,
                author=extracted_metadata.get("author"),
                page_count=extracted_metadata.get("page_count"),
                word_count=extracted_metadata.get("word_count"),
                file_size=file_path.stat().st_size,
                file_type=file_type,
                tags=tags or [],
                language=extracted_metadata.get("language"),
            )

            extraction_time = time.time() - start_time

            logger.info(
                f"Document processed successfully: {document_id} "
                f"({extraction_time:.2f}s)"
            )

            return DocumentExtractionResult(
                document_id=document_id,
                filename=filename,
                file_type=file_type,
                text_content=text_content,
                metadata=metadata,
                extraction_time=extraction_time,
                success=True,
                error_message=None,
            )

        except Exception as e:
            extraction_time = time.time() - start_time
            logger.error(f"Document processing failed: {e}")
            filename = file.filename or "unknown"

            return DocumentExtractionResult(
                document_id="",
                filename=filename,
                file_type=DocumentType.PDF,
                text_content="",
                metadata=DocumentMetadata(
                    title=None,
                    author=None,
                    page_count=None,
                    word_count=None,
                    file_size=0,
                    file_type=DocumentType.PDF,
                    tags=[],
                    language=None,
                ),
                extraction_time=extraction_time,
                success=False,
                error_message=str(e),
            )

# Global instance
_document_processor: Optional[DocumentProcessor] = None

def get_document_processor() -> DocumentProcessor:
    """Get or create document processor instance."""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor